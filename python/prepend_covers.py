"""Prepend cover images to 家谱.docx.

The script takes 16 cover images (``1.jpeg`` ... ``16.jpeg``) from a folder
and produces a new combined Word document where:

* The first 16 pages are the cover images (one per page, no extra text).
* The original 431 genealogy pages follow afterwards, completely untouched.

Implementation outline
----------------------
1. Build a temporary ``封面.docx`` in-memory:
   * page size & margins are copied from the existing 家谱.docx so covers
     match the rest of the book exactly;
   * each page contains a single image scaled to fill the available area
     while preserving its aspect ratio;
   * a hard page break is placed after every image except the last.
2. Use :class:`docxcompose.composer.Composer` to append the original
   家谱.docx after the cover document.  ``docxcompose`` keeps every
   appended section's ``sectPr`` intact, so the original 431 pages are
   bit-for-bit identical to before.
3. Save the merged file back to 家谱.docx (the original is overwritten in
   place after a backup copy).
"""

from __future__ import annotations

import io
import shutil
import sys
import time
from pathlib import Path
from typing import List, Tuple

from PIL import Image
from docx import Document
from docx.shared import Emu
from docx.enum.text import WD_BREAK
from docxcompose.composer import Composer


COVER_DIR = Path(r"C:\Users\alexdu\Downloads\vFlat\封面")
GENEALOGY_DOCX = Path(r"C:\Users\alexdu\Downloads\vFlat\家谱.docx")
COVER_COUNT = 16
EMU_PER_INCH = 914400


def collect_covers(src: Path, count: int) -> List[Path]:
    """Return ``[1.jpeg, 2.jpeg, ..., count.jpeg]`` (any common ext)."""

    exts = (".jpeg", ".jpg", ".png")
    files: List[Path] = []
    for i in range(1, count + 1):
        match: Path | None = None
        for ext in exts:
            cand = src / f"{i}{ext}"
            if cand.is_file():
                match = cand
                break
        if match is None:
            raise FileNotFoundError(f"Missing cover image #{i} in {src}")
        files.append(match)
    return files


def fit_size(img_path: Path, max_w_emu: int, max_h_emu: int) -> Tuple[int, int]:
    """Compute (width_emu, height_emu) so the image fits the box while
    preserving its aspect ratio."""

    with Image.open(img_path) as im:
        w_px, h_px = im.size
    if w_px <= 0 or h_px <= 0:
        raise ValueError(f"Invalid image dimensions for {img_path}")
    ratio = w_px / h_px
    box_ratio = max_w_emu / max_h_emu
    if ratio >= box_ratio:
        # Image is wider relative to the box -> width is the limiting side.
        w_emu = max_w_emu
        h_emu = int(round(max_w_emu / ratio))
    else:
        h_emu = max_h_emu
        w_emu = int(round(max_h_emu * ratio))
    return w_emu, h_emu


def build_cover_doc(covers: List[Path], reference_docx: Path) -> Document:
    """Create an in-memory Document containing one cover image per page."""

    # Start from a *copy* of the reference doc so we inherit its section
    # properties (page size, margins) exactly, then strip its body content.
    doc = Document(str(reference_docx))
    body = doc.element.body
    # Remove every block-level element except the trailing sectPr.
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)

    section = doc.sections[0]
    page_w = section.page_width
    page_h = section.page_height
    avail_w_emu = page_w - section.left_margin - section.right_margin
    avail_h_emu = page_h - section.top_margin - section.bottom_margin
    print(
        f"[info] cover page box: {avail_w_emu / EMU_PER_INCH:.3f}\" x "
        f"{avail_h_emu / EMU_PER_INCH:.3f}\""
    )

    for idx, img_path in enumerate(covers, start=1):
        w_emu, h_emu = fit_size(img_path, avail_w_emu, avail_h_emu)
        para = doc.add_paragraph()
        # Tight paragraph so the image isn't pushed to a second page by
        # paragraph spacing.
        pf = para.paragraph_format
        pf.space_before = Emu(0)
        pf.space_after = Emu(0)
        run = para.add_run()
        run.add_picture(str(img_path), width=Emu(w_emu), height=Emu(h_emu))
        # Add a hard page break after every cover except the last; the last
        # one is followed by the original 家谱 content (which itself starts
        # a new section) so no extra break is needed.
        if idx < len(covers):
            run.add_break(WD_BREAK.PAGE)
        print(
            f"[info] cover {idx}/{len(covers)} -> {img_path.name} "
            f"({w_emu / EMU_PER_INCH:.2f}\" x {h_emu / EMU_PER_INCH:.2f}\")"
        )

    return doc


def main() -> None:
    if not GENEALOGY_DOCX.is_file():
        raise SystemExit(f"Genealogy docx not found: {GENEALOGY_DOCX}")
    if not COVER_DIR.is_dir():
        raise SystemExit(f"Cover folder not found: {COVER_DIR}")

    covers = collect_covers(COVER_DIR, COVER_COUNT)
    print(f"[info] collected {len(covers)} cover images")

    print("[info] building cover document ...")
    cover_doc = build_cover_doc(covers, GENEALOGY_DOCX)

    # Persist to a buffer, then load again as the *base* of the final merge.
    # docxcompose works most reliably when the base is a fresh Document
    # loaded from disk/bytes.
    print("[info] serialising cover document to memory ...")
    buf = io.BytesIO()
    cover_doc.save(buf)
    buf.seek(0)

    print("[info] merging cover + 家谱 ...")
    base = Document(buf)
    composer = Composer(base)
    started = time.time()
    composer.append(Document(str(GENEALOGY_DOCX)))
    print(f"[info] append done in {time.time() - started:.1f}s")

    # Backup the original then overwrite in place.
    backup = GENEALOGY_DOCX.with_name(GENEALOGY_DOCX.stem + "_no_cover" + GENEALOGY_DOCX.suffix)
    if not backup.exists():
        print(f"[info] backing up original -> {backup.name}")
        shutil.copy2(GENEALOGY_DOCX, backup)
    else:
        print(f"[info] backup already exists, leaving it untouched: {backup.name}")

    print(f"[info] saving merged file -> {GENEALOGY_DOCX}")
    composer.save(str(GENEALOGY_DOCX))
    size_mb = GENEALOGY_DOCX.stat().st_size / 1024 / 1024
    print(f"[done] wrote {GENEALOGY_DOCX} ({size_mb:.2f} MB)")


if __name__ == "__main__":
    try:
        main()
    except Exception as exc:  # pragma: no cover - top-level guard
        print(f"[fatal] {exc}", file=sys.stderr)
        raise
