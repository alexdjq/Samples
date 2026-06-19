# -*- coding: utf-8 -*-
"""
text_json_to_word.py
====================

Convert a PP-StructureV3 JSON file produced by ``paddleocr_cli.py`` into a
Word ``.docx`` file for **text-heavy** genealogy pages (prefaces, narratives,
single-column body text -- as opposed to the multi-column table pages that
``json_to_word.py`` handles).

The output mimics the printed page: an A4 sheet whose body is wrapped in a
single rectangular frame ("外框"), with a top-level title, an optional
section heading and the body paragraphs flowing inside.

Usage
-----
Single file:
    python text_json_to_word.py <input.json> [output.docx]

Whole directory (every ``*.json`` next to a ``*.jpg`` becomes ``*.docx``):
    python text_json_to_word.py -d <dir> [--overwrite]

Dependencies:
    pip install python-docx
"""

from __future__ import annotations

import argparse
import json
import os
import re
import sys
from typing import List, Tuple, Dict, Any, Optional

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Missing dependency 'python-docx'. Please install it via:")
    print("    pip install python-docx")
    sys.exit(1)


# ---------------------------------------------------------------------------
# Page / frame layout constants -- kept in sync with json_to_word.py so the
# table pages and text pages share the same printed dimensions.
# ---------------------------------------------------------------------------

A4_WIDTH_CM  = 21.0
A4_HEIGHT_CM = 29.7

MARGIN_TOP_CM    = 0.5
MARGIN_BOTTOM_CM = 0.5
MARGIN_LEFT_CM   = 1.0
MARGIN_RIGHT_CM  = 1.0

# Body typography
TITLE_FONT_PT    = 18.0   # ## 渑池杜氏宗谱  (now rendered in the spine column)
HEADING_FONT_PT  = 16.0   # # 重修杜氏宗谱序
BODY_FONT_PT     = 12.0   # default body size; auto-shrunk when needed
BODY_FONT_PT_MIN = 8.0    # never shrink below this -- keeps text legible
LINE_SPACING     = 1.4    # default line spacing for long paragraphs
LINE_SPACING_MIN = 1.05   # collapse line spacing alongside the font size
PARA_INDENT_CHARS = 2     # first-line indent: 2 chinese chars

# Two distinct CJK faces: KaiTi for the flowing body (calligraphic look
# requested by the user) and SimSun as a robust fallback for headings,
# the spine column and any environment that lacks KaiTi.
BODY_FONT    = "KaiTi"
CHINESE_FONT = "SimSun"

# ---------------------------------------------------------------------------
# Right-hand "spine" column (mirrors the merged header column on the
# table pages).  Width, font face and weight are deliberately matched to
# the spine column in ``json_to_word.py`` so the table-page and text-page
# Word documents share the *exact* same spine appearance.
#
#  * Width: the same per-cell width used by the 24-column genealogy
#    table -- ``(usable_w) / 24`` cm.  This is noticeably narrower
#    (~0.79 cm) than the previous 1.2 cm, matching the printed PDF.
#  * Font:  SimSun, **bold** -- exactly what ``set_cell_text`` writes
#    for the merged title/page cell on the table pages.
#  * Size:  12 pt -- the upper bound used by ``_choose_font_size`` in
#    json_to_word.py for the merged spine cell on table pages.
# ---------------------------------------------------------------------------
SPINE_COL_WIDTH_CM = (A4_WIDTH_CM - MARGIN_LEFT_CM - MARGIN_RIGHT_CM) / 24.0
SPINE_FONT_PT      = 12.0
SPINE_TEXT         = "渑池  杜氏宗谱"   # two spaces between the two parts

# ---------------------------------------------------------------------------
# Image-only page whitelist
# ---------------------------------------------------------------------------
# Some source pages are predominantly photographic / pictorial (portraits,
# scanned illustrations, ...) and OCR text is meaningless for them.  For
# every page number listed below the body region is filled with the
# original scan image (``<stem>.jpg``) instead of the parsed markdown,
# while the outer frame and right-hand spine column are kept identical
# to the regular text pages so the bound book stays visually consistent
# and the page numbers remain continuous.
IMAGE_PAGE_WHITELIST: "frozenset[int]" = frozenset({24, 25, 26, 65, 431})

# Acceptable source-image extensions, tried in order.  The first one that
# exists next to the JSON file wins.
_IMAGE_EXTENSIONS = (".jpg", ".jpeg", ".png", ".JPG", ".JPEG", ".PNG")

# ---------------------------------------------------------------------------
# Page-number helpers -- the spine column ends with a Chinese-numeral page
# label derived from the trailing integer in the source filename, e.g.
# ``Default folder - 3.json`` -> ``三``.
# ---------------------------------------------------------------------------

_CN_DIGITS = "零一二三四五六七八九"
_CN_UNITS_SMALL = ["", "十", "百", "千"]
_CN_UNITS_BIG   = ["", "万", "亿"]

def _int_to_chinese_numeral(n: int) -> str:
    """Convert a non-negative integer to its traditional Chinese numeral
    form (e.g. 3 -> '三', 11 -> '十一', 234 -> '二百三十四').

    The conversion follows the standard rules:
    * leading ``1`` in the tens place is dropped (``11`` -> ``十一``);
    * runs of zero digits collapse to a single ``零``;
    * a trailing zero never produces a stray ``零``.
    """
    if n < 0:
        raise ValueError("page number must be non-negative")
    if n == 0:
        return _CN_DIGITS[0]
    if n < 10:
        return _CN_DIGITS[n]
    if n < 20:
        # 10 -> '十', 11 -> '十一', ..., 19 -> '十九'
        return "十" + ("" if n == 10 else _CN_DIGITS[n - 10])

    # General case: split into groups of four digits (亿, 万, base).
    groups: List[int] = []
    rem = n
    while rem > 0:
        groups.append(rem % 10000)
        rem //= 10000

    parts: List[str] = []
    for idx in range(len(groups) - 1, -1, -1):
        g = groups[idx]
        if g == 0:
            # Zero group -- emit a single '零' to preserve readability,
            # but avoid duplicating it.
            if parts and not parts[-1].endswith(_CN_DIGITS[0]):
                parts.append(_CN_DIGITS[0])
            continue
        # Convert this 4-digit group.
        digits = [(g // 1000) % 10, (g // 100) % 10, (g // 10) % 10, g % 10]
        chunk: List[str] = []
        zero_pending = False
        for pos, d in enumerate(digits):
            unit = _CN_UNITS_SMALL[3 - pos]
            if d == 0:
                zero_pending = True
            else:
                if zero_pending and chunk:
                    chunk.append(_CN_DIGITS[0])
                chunk.append(_CN_DIGITS[d] + unit)
                zero_pending = False
        parts.append("".join(chunk) + _CN_UNITS_BIG[idx])

    result = "".join(parts)
    # Tidy up: collapse runs of '零' and strip a trailing '零'.
    while "零零" in result:
        result = result.replace("零零", "零")
    if result.endswith("零"):
        result = result[:-1]
    return result

def _extract_page_label(input_path: str) -> Optional[str]:
    """Pull the trailing integer out of ``input_path`` and convert it to a
    Chinese numeral suitable for use as a page label.  Returns ``None``
    when no trailing number can be found.
    """
    stem = os.path.splitext(os.path.basename(input_path))[0]
    m = re.search(r"(\d+)\s*$", stem)
    if not m:
        return None
    try:
        return _int_to_chinese_numeral(int(m.group(1)))
    except ValueError:
        return None


def _extract_page_number(input_path: str) -> Optional[int]:
    """Return the trailing integer in ``input_path`` (e.g. ``24``) or
    ``None`` if the filename has no trailing number.  This is the
    integer counterpart to :func:`_extract_page_label`, used to test
    membership in :data:`IMAGE_PAGE_WHITELIST`.
    """
    stem = os.path.splitext(os.path.basename(input_path))[0]
    m = re.search(r"(\d+)\s*$", stem)
    if not m:
        return None
    try:
        return int(m.group(1))
    except ValueError:
        return None


def _find_source_image(json_path: str) -> Optional[str]:
    """Locate the original scan image that sits next to ``json_path``.

    Tries every extension in :data:`_IMAGE_EXTENSIONS` in order and
    returns the first existing file's absolute path.  Returns ``None``
    when no companion image can be found.
    """
    base, _ = os.path.splitext(os.path.abspath(json_path))
    for ext in _IMAGE_EXTENSIONS:
        candidate = base + ext
        if os.path.isfile(candidate):
            return candidate
    return None


# ---------------------------------------------------------------------------
# Markdown parsing helpers
# ---------------------------------------------------------------------------

# A "block" describes one logical piece of content extracted from the
# OCR markdown.  ``kind`` is one of: ``"title"``, ``"heading"``,
# ``"subheading"`` or ``"paragraph"``.
Block = Tuple[str, str]


def _parse_markdown_blocks(md_text: str) -> List[Block]:
    """Split the OCR markdown text into a list of typed blocks.

    The PP-StructureV3 markdown output for a preface page typically looks
    like::

        ## 渑池杜氏宗谱

        # 重修杜氏宗谱序

        宗谱，又称家谱、家乘 ...

        吾族系出尧帝 ...

    We honour the heading hierarchy: ``##`` becomes the document title
    (printed once at the top), ``#`` becomes a centred section heading,
    everything else is treated as a body paragraph.  Blank lines act as
    paragraph separators.
    """
    blocks: List[Block] = []
    # Normalise newlines and collapse hard-wrapped lines: PP-Structure
    # often inserts a newline every ~30-40 chars even though the source
    # paragraph is one continuous block.  We keep blank lines (paragraph
    # separators) but glue any other adjacent non-heading lines back
    # together.
    raw_lines = md_text.replace("\r\n", "\n").replace("\r", "\n").split("\n")

    buf: List[str] = []

    def _flush_paragraph() -> None:
        if not buf:
            return
        # Glue without spaces -- Chinese text has no inter-word space.
        joined = "".join(s.strip() for s in buf).strip()
        if joined:
            blocks.append(("paragraph", joined))
        buf.clear()

    for line in raw_lines:
        stripped = line.strip()
        if not stripped:
            _flush_paragraph()
            continue
        if stripped.startswith("## "):
            _flush_paragraph()
            blocks.append(("title", stripped[3:].strip()))
            continue
        if stripped.startswith("# "):
            _flush_paragraph()
            blocks.append(("heading", stripped[2:].strip()))
            continue
        if stripped.startswith("### "):
            _flush_paragraph()
            blocks.append(("subheading", stripped[4:].strip()))
            continue
        buf.append(stripped)

    _flush_paragraph()
    return blocks


# ---------------------------------------------------------------------------
# Word helpers (mirrors the conventions used in json_to_word.py)
# ---------------------------------------------------------------------------

def _setup_a4_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width    = Cm(A4_WIDTH_CM)
    section.page_height   = Cm(A4_HEIGHT_CM)
    section.top_margin    = Cm(MARGIN_TOP_CM)
    section.bottom_margin = Cm(MARGIN_BOTTOM_CM)
    section.left_margin   = Cm(MARGIN_LEFT_CM)
    section.right_margin  = Cm(MARGIN_RIGHT_CM)


def _set_run_chinese_font(run, font_name: str = CHINESE_FONT) -> None:
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)


def _set_cell_borders(cell,
                      top: str = "single", bottom: str = "single",
                      left: str = "single", right: str = "single",
                      sz: str = "8") -> None:
    """Set per-cell borders.  ``sz`` is in eighths of a point."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for side, val in (("top", top), ("left", left),
                      ("bottom", bottom), ("right", right)):
        node = tcBorders.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tcBorders.append(node)
        node.set(qn("w:val"),   val)
        node.set(qn("w:sz"),    sz)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "auto")


def _set_table_fixed_layout(table) -> None:
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    layout = tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")


def _set_cell_margins(cell,
                      top_cm: float = 0.4, bottom_cm: float = 0.4,
                      left_cm: float = 0.5, right_cm: float = 0.5) -> None:
    """Inner padding so the body text doesn't touch the outer frame."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    mar = tcPr.find(qn("w:tcMar"))
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tcPr.append(mar)

    def _to_dxa(cm: float) -> str:
        # 1 cm == 567 twentieths of a point (dxa)
        return str(int(round(cm * 567)))

    for side, cm in (("top", top_cm), ("left", left_cm),
                     ("bottom", bottom_cm), ("right", right_cm)):
        node = mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            mar.append(node)
        node.set(qn("w:w"),    _to_dxa(cm))
        node.set(qn("w:type"), "dxa")


def _set_cell_text_direction(cell, direction: str = "tbRl") -> None:
    """Make the cell's text run top-to-bottom (vertical CJK style).

    ``tbRl`` (top-to-bottom, right-to-left) is what Word uses for
    classical vertical Chinese typesetting -- characters stack
    downwards and successive lines stack right-to-left.  Combined with
    a centred paragraph this gives the "book spine" look.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    td = tcPr.find(qn("w:textDirection"))
    if td is None:
        td = OxmlElement("w:textDirection")
        tcPr.append(td)
    td.set(qn("w:val"), direction)


def _set_row_exact_height(row, height_cm: float) -> None:
    """Force the table row to render at *exactly* ``height_cm`` cm.

    Word treats the height attribute on a ``<w:tr>`` as a *minimum* by
    default, which means a row with little content collapses regardless
    of the requested height.  Setting ``hRule="exact"`` pins the row so
    the outer frame keeps a constant size whether the body has many
    paragraphs or a single short one.  This is exactly what the user
    asks for: the frame should look identical no matter how small the
    body font ended up.
    """
    tr = row._tr
    trPr = tr.find(qn("w:trPr"))
    if trPr is None:
        trPr = OxmlElement("w:trPr")
        tr.insert(0, trPr)
    height_node = trPr.find(qn("w:trHeight"))
    if height_node is None:
        height_node = OxmlElement("w:trHeight")
        trPr.append(height_node)
    # 1 cm == 567 dxa (twentieths of a point).
    height_node.set(qn("w:val"), str(int(round(height_cm * 567))))
    height_node.set(qn("w:hRule"), "exact")
    # Also forbid the row from splitting across pages -- the outer
    # frame must stay together as one rectangle.
    cant_split = trPr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        trPr.append(cant_split)




# ---------------------------------------------------------------------------
# Block rendering
# ---------------------------------------------------------------------------

def _add_run(p, text: str, *, size_pt: float, bold: bool = False,
             font_name: str = CHINESE_FONT) -> None:
    run = p.add_run(text)
    run.font.size = Pt(size_pt)
    run.font.name = font_name
    run.bold = bold
    _set_run_chinese_font(run, font_name)


def _render_spine_cell(cell, text: str = SPINE_TEXT,
                       page_label: Optional[str] = None) -> None:
    """Render the book-title spine column on the right edge.

    Each Chinese character is placed on its own centred paragraph so
    that all glyphs stand upright while stacking from top to bottom
    -- the classical Chinese vertical layout.  Spaces in ``text`` are
    rendered as empty paragraphs, giving the small visual gap that
    separates ``渑池`` from ``杜氏宗谱``.

    When ``page_label`` is supplied (e.g. ``"三"``), it is appended
    after ``杜氏宗谱`` with a single visual gap, so the spine reads
    top-to-bottom as ``渑池   杜氏宗谱   三`` -- mirroring the
    printed page.
    """
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # First character reuses the auto-inserted paragraph; subsequent
    # ones are appended one per character.
    first_para = cell.paragraphs[0]
    first_used = False

    def _new_para():
        nonlocal first_used
        if not first_used:
            first_used = True
            return first_para
        return cell.add_paragraph()

    full_text = text
    if page_label:
        # One space-gap between '杜氏宗谱' and the page label.
        full_text = f"{text}  {page_label}"

    for ch in full_text:
        p = _new_para()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after  = Pt(0)
        pf.line_spacing = 1.0
        if ch == " ":
            # Empty paragraph -> visual gap between segments.
            continue
        # Match the spine glyph style used in json_to_word.py:
        # SimSun face, bold weight -- so the text-page and table-page
        # spine columns are visually indistinguishable.
        _add_run(p, ch, size_pt=SPINE_FONT_PT, bold=True,
                 font_name=CHINESE_FONT)


def _estimate_body_font_pt(blocks: List[Block],
                           body_w_cm: float,
                           body_h_cm: float) -> Tuple[float, float]:
    """Pick a body font size (and matching line spacing) that should
    keep the rendered document within a single A4 page.

    Strategy: prefer the largest legible font.  For each candidate
    font size (descending from :data:`BODY_FONT_PT`), try the full
    line-spacing range from :data:`LINE_SPACING` down to
    :data:`LINE_SPACING_MIN` and accept the first combination whose
    estimated rendered height fits the available cell area.  This
    means we'll happily tighten line spacing before shrinking the
    font, which is what produces the best visual result on a
    single-page layout.

    Returns a tuple ``(body_font_pt, line_spacing)``.
    """
    # Convert cell area from centimetres to points (1 cm == 28.3465 pt).
    cm_to_pt = 28.3465
    cell_w_pt = body_w_cm * cm_to_pt
    cell_h_pt = body_h_cm * cm_to_pt

    body_chars = 0
    heading_count    = 0
    subheading_count = 0
    paragraph_count  = 0
    for kind, text in blocks:
        if kind == "title":
            continue
        if kind == "heading":
            heading_count += 1
        elif kind == "subheading":
            subheading_count += 1
        else:
            body_chars += len(text)
            paragraph_count += 1

    # Walk font size from preferred down to the floor; for each size
    # walk line spacing from preferred down to the tight floor.
    size = BODY_FONT_PT
    while size >= BODY_FONT_PT_MIN:
        scale = size / BODY_FONT_PT
        heading_h_pt    = HEADING_FONT_PT     * scale * 1.3 + 18 * scale
        subheading_h_pt = (BODY_FONT_PT + 1)  * scale * 1.3 + 12 * scale
        # Each body paragraph adds a small ``space_after`` of 3pt*scale
        # plus its first-line indent eats two character slots on the
        # opening line.
        per_para_extra_chars = PARA_INDENT_CHARS
        per_para_extra_pt    = 3 * scale

        # Characters that fit on one line of body text.  Use a slightly
        # generous divisor (0.98) because Chinese glyphs are nominally
        # square but Word also fits punctuation tighter at the line end.
        chars_per_line = max(1, int(cell_w_pt / (size * 0.98)))

        spacing = LINE_SPACING
        while spacing >= LINE_SPACING_MIN - 1e-9:
            # Body lines: account for the per-paragraph indent.
            effective_chars = body_chars + paragraph_count * per_para_extra_chars
            body_lines = -(-effective_chars // chars_per_line)  # ceil
            line_height_pt = size * spacing
            needed_h = (
                body_lines * line_height_pt
                + paragraph_count * per_para_extra_pt
                + heading_count    * heading_h_pt
                + subheading_count * subheading_h_pt
            )
            if needed_h <= cell_h_pt:
                return size, spacing
            spacing -= 0.05

        size -= 0.5

    # Fell off the bottom -- return the floor so the caller still has
    # a deterministic value to work with.
    return BODY_FONT_PT_MIN, LINE_SPACING_MIN


def _render_blocks_into_cell(cell, blocks: List[Block],
                             body_font_pt: float = BODY_FONT_PT,
                             line_spacing: float = LINE_SPACING) -> None:
    """Write the parsed blocks into the body (left) cell.

    The ``title`` blocks (``## 渑池杜氏宗谱``) are intentionally skipped
    here -- they are rendered separately in the right-hand spine
    column, just like the merged header column on the table pages.

    ``body_font_pt`` and ``line_spacing`` are the values picked by
    :func:`_estimate_body_font_pt` so the page stays within a single
    A4 sheet.  Heading sizes scale proportionally so the visual
    hierarchy is preserved when the body has been shrunk.
    """
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # Scale headings in proportion to the (possibly shrunk) body size.
    scale = body_font_pt / BODY_FONT_PT
    heading_pt    = HEADING_FONT_PT * scale
    subheading_pt = (BODY_FONT_PT + 1) * scale

    # python-docx pre-creates one empty paragraph in every cell -- reuse
    # it for the very first block so we don't end up with a leading blank.
    first_para = cell.paragraphs[0]
    first_used = False

    def _new_para():
        nonlocal first_used
        if not first_used:
            first_used = True
            return first_para
        return cell.add_paragraph()

    for kind, text in blocks:
        if kind == "title":
            # Rendered in the spine column; do not duplicate here.
            continue

        p = _new_para()
        pf = p.paragraph_format
        pf.line_spacing = line_spacing

        if kind == "heading":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.space_before = Pt(6 * scale)
            pf.space_after  = Pt(12 * scale)
            # Headings use SimSun (bold) for solid emphasis.
            _add_run(p, text, size_pt=heading_pt, bold=True,
                     font_name=CHINESE_FONT)
        elif kind == "subheading":
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.space_before = Pt(6 * scale)
            pf.space_after  = Pt(6 * scale)
            _add_run(p, text, size_pt=subheading_pt, bold=True,
                     font_name=CHINESE_FONT)
        else:  # paragraph -- KaiTi as requested by the user.
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf.space_before = Pt(0)
            pf.space_after  = Pt(3 * scale)
            # 2 Chinese-character first-line indent.
            pf.first_line_indent = Pt(body_font_pt * PARA_INDENT_CHARS)
            _add_run(p, text, size_pt=body_font_pt, bold=False,
                     font_name=BODY_FONT)


def _render_image_into_cell(cell, image_path: str,
                            cell_w_cm: float, cell_h_cm: float) -> None:
    """Replace the body cell's content with the original scan image.

    The picture is centred both horizontally and vertically inside the
    cell, with its dimensions clamped so it never overflows.  Aspect
    ratio is preserved by sizing only the dominant dimension (the one
    that hits the cell limit first); python-docx then computes the
    other side automatically.
    """
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Reuse the auto-inserted first paragraph and drop any others -- a
    # blank trailing paragraph would push the image upward inside the
    # exact-height row and break the centring.
    paragraphs = list(cell.paragraphs)
    p = paragraphs[0]
    for extra in paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    # Clear any pre-existing runs in the first paragraph.
    for run_elem in list(p._element.findall(qn("w:r"))):
        p._element.remove(run_elem)

    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)
    pf.line_spacing = 1.0

    # Leave a small inner margin so the image never visually touches
    # the cell border.  Cell padding is already 0.3 cm horizontally /
    # 0.3 cm vertically (see convert_json_to_docx); subtract a touch
    # more to keep things airy.
    max_w_cm = max(0.5, cell_w_cm - 0.4)
    max_h_cm = max(0.5, cell_h_cm - 0.4)

    # Decide which dimension to constrain by inspecting the source
    # image's aspect ratio.  python-docx requires a width *or* a height
    # (not both) to preserve the natural ratio.
    try:
        from PIL import Image  # type: ignore  # optional dependency
        with Image.open(image_path) as im:
            img_w_px, img_h_px = im.size
    except Exception:                              # noqa: BLE001
        img_w_px = img_h_px = 0

    run = p.add_run()
    if img_w_px > 0 and img_h_px > 0:
        # Cell aspect = w/h.  If the image is *wider* relative to the
        # cell, width is the binding constraint; otherwise height is.
        cell_ratio = max_w_cm / max_h_cm
        img_ratio  = img_w_px / img_h_px
        if img_ratio >= cell_ratio:
            run.add_picture(image_path, width=Cm(max_w_cm))
        else:
            run.add_picture(image_path, height=Cm(max_h_cm))
    else:
        # Pillow unavailable or unreadable image -- fall back to
        # constraining the width only; Word may letterbox vertically
        # but at least nothing overflows the page.
        run.add_picture(image_path, width=Cm(max_w_cm))


# ---------------------------------------------------------------------------
# Top-level conversion
# ---------------------------------------------------------------------------

def _extract_markdown_text(json_obj: Dict[str, Any]) -> str:
    """Best-effort extraction of the raw markdown text from a PP-Structure
    JSON file.  Both the per-page payload and the per-file payload that
    ``paddleocr_cli.py`` saves are accepted.

    A *present-but-empty* ``markdown.text`` (which PP-StructureV3 emits
    for blank source pages) is treated as a successful extraction of the
    empty string -- callers downstream are responsible for rendering an
    appropriate placeholder.  Only payloads that genuinely lack the
    field altogether raise ``ValueError``."""
    # Per-file payload as saved by paddleocr_cli.py
    lpr = json_obj.get("layoutParsingResults")
    if isinstance(lpr, list) and lpr:
        saw_field = False
        chunks: List[str] = []
        for entry in lpr:
            md = entry.get("markdown") if isinstance(entry, dict) else None
            if isinstance(md, dict) and "text" in md and isinstance(md["text"], str):
                saw_field = True
                if md["text"]:
                    chunks.append(md["text"])
        if saw_field:
            return "\n\n".join(chunks)

    # Some payloads expose ``markdown`` at the top level.
    md = json_obj.get("markdown")
    if isinstance(md, dict) and "text" in md and isinstance(md["text"], str):
        return md["text"]
    if isinstance(md, str):
        return md

    raise ValueError(
        "Could not find a 'markdown.text' field in the JSON payload. "
        "Is this really a PP-StructureV3 result?"
    )


def convert_json_to_docx(input_json: str, output_docx: str) -> None:
    with open(input_json, "r", encoding="utf-8") as f:
        payload = json.load(f)

    # Image-only pages bypass markdown parsing entirely -- the body
    # cell is filled with the original scan instead.
    page_no = _extract_page_number(input_json)
    is_image_page = page_no is not None and page_no in IMAGE_PAGE_WHITELIST
    image_path: Optional[str] = None
    if is_image_page:
        image_path = _find_source_image(input_json)
        if image_path is None:
            raise FileNotFoundError(
                f"Page {page_no} is marked as an image page but no "
                f"companion image (.jpg/.jpeg/.png) was found next to "
                f"{input_json!r}."
            )

    if is_image_page:
        # Skip markdown extraction altogether; image pages do not need
        # any blocks.
        md_text = ""
        blocks: List[Block] = []
    else:
        md_text = _extract_markdown_text(payload)
        blocks = _parse_markdown_blocks(md_text)
    # Blank source pages (no OCR text whatsoever) still need to produce
    # a physical .docx so the printed page numbering stays continuous
    # with the surrounding pages.  We render the same outer frame +
    # spine column, but leave the body cell empty.
    is_blank_page = (not is_image_page) and (not blocks)

    doc = Document()
    _setup_a4_page(doc)

    # Decide the body font size up-front so we can drive both the
    # Normal style and the per-paragraph runs from a single source of
    # truth.
    usable_w_cm = A4_WIDTH_CM  - MARGIN_LEFT_CM - MARGIN_RIGHT_CM
    usable_h_cm = A4_HEIGHT_CM - MARGIN_TOP_CM  - MARGIN_BOTTOM_CM
    body_w_cm  = usable_w_cm - SPINE_COL_WIDTH_CM
    # Subtract the cell padding (0.5 + 0.5 cm horizontal, 0.4 + 0.4 cm
    # vertical) so the estimate matches the actually-printable region.
    if is_blank_page or is_image_page:
        # No content to size against -- pick the same defaults as a
        # populated page so the spine column's font matches its peers.
        body_pt, body_spacing = 12.0, 1.5
    else:
        body_pt, body_spacing = _estimate_body_font_pt(
            blocks,
            body_w_cm=body_w_cm - 0.6,   # minus 0.3 cm left + 0.3 cm right padding
            body_h_cm=usable_h_cm - 0.6, # minus 0.3 cm top  + 0.3 cm bottom padding
        )

    # Tighten the Normal style so cell paragraphs control spacing.
    try:
        normal = doc.styles["Normal"]
        normal.font.size = Pt(body_pt)
        normal.font.name = BODY_FONT
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after  = Pt(0)
        normal.paragraph_format.line_spacing = body_spacing
    except KeyError:
        pass

    # Strip the auto-inserted leading paragraph so the frame sits right
    # at the top of the page.
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)

    # ---- Outer frame -------------------------------------------------
    # Regular pages get a 1x2 table (body + right-hand spine column).
    # Image-only pages drop the spine column entirely -- the embedded
    # scan already carries the title, page number and frame -- so we
    # build a 1x1 table whose single cell spans the full body width.
    if is_image_page:
        table = doc.add_table(rows=1, cols=1)
        _set_table_fixed_layout(table)
        # Single column spans the entire usable width.
        full_w_cm = body_w_cm + SPINE_COL_WIDTH_CM
        table.columns[0].width = Cm(full_w_cm)
    else:
        table = doc.add_table(rows=1, cols=2)
        _set_table_fixed_layout(table)
        table.columns[0].width = Cm(body_w_cm)
        table.columns[1].width = Cm(SPINE_COL_WIDTH_CM)

    # Pin the outer frame's height so it always fills the printable
    # area, regardless of how short the rendered body text turns out
    # to be.  Without ``hRule="exact"`` Word treats the row height as
    # a minimum and lets the frame collapse around the content,
    # producing a stunted-looking box when the body fits in just a
    # few lines.
    _set_row_exact_height(table.rows[0], usable_h_cm)

    cell = table.cell(0, 0)
    if is_image_page:
        # Body cell spans the full usable width on image pages.
        cell.width  = Cm(body_w_cm + SPINE_COL_WIDTH_CM)
    else:
        cell.width  = Cm(body_w_cm)
    cell.height = Cm(usable_h_cm)
    # Slightly tighter horizontal padding so text uses more of the
    # available width and avoids overflowing onto a second page.
    _set_cell_margins(cell,
                      top_cm=0.3, bottom_cm=0.3,
                      left_cm=0.3, right_cm=0.3)
    # Image-only pages ship the frame inside the embedded scan itself,
    # so we suppress *all* borders for them to avoid drawing a redundant
    # outer line.
    _border_val = "nil" if is_image_page else "single"
    _set_cell_borders(cell,
                      top=_border_val, bottom=_border_val,
                      left=_border_val, right=_border_val, sz="8")

    # The spine column only exists on regular (non-image) pages.
    spine_cell = None
    if not is_image_page:
        spine_cell = table.cell(0, 1)
        spine_cell.width  = Cm(SPINE_COL_WIDTH_CM)
        spine_cell.height = Cm(usable_h_cm)
        # Tight padding on the narrow spine column so the characters fit.
        _set_cell_margins(spine_cell,
                          top_cm=0.2, bottom_cm=0.2,
                          left_cm=0.05, right_cm=0.05)
        _set_cell_borders(spine_cell,
                          top=_border_val, bottom=_border_val,
                          left=_border_val, right=_border_val, sz="8")

    if is_image_page:
        # Whitelisted image-only page: drop the original scan into the
        # body cell, preserving aspect ratio and centring it inside
        # the fixed-height frame.  No spine column / page label is
        # rendered -- the scan already includes them.
        assert image_path is not None  # narrowed above
        full_w_cm = body_w_cm + SPINE_COL_WIDTH_CM
        _render_image_into_cell(cell, image_path,
                                cell_w_cm=full_w_cm,
                                cell_h_cm=usable_h_cm)
    elif is_blank_page:
        # Insert one empty paragraph so the cell occupies its full
        # exact-height row instead of collapsing.  The paragraph has
        # no runs, so nothing prints in the body region.
        cell.paragraphs[0].text = ""
    else:
        _render_blocks_into_cell(cell, blocks,
                                 body_font_pt=body_pt,
                                 line_spacing=body_spacing)

    # Spine column is only present on regular pages.
    if spine_cell is not None:
        page_label = _extract_page_label(input_json)
        _render_spine_cell(spine_cell, page_label=page_label)

    # Make sure the output directory exists.
    out_dir = os.path.dirname(os.path.abspath(output_docx))
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)
    doc.save(output_docx)


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def _is_text_page(payload: Dict[str, Any]) -> bool:
    """Heuristic: a page is "text-heavy" when its markdown contains no
    HTML/markdown table at all.  This lets the directory mode skip the
    table pages (which json_to_word.py handles) automatically."""
    try:
        md_text = _extract_markdown_text(payload)
    except ValueError:
        return False
    lowered = md_text.lower()
    if "<table" in lowered or "</table>" in lowered:
        return False
    # A markdown pipe-table needs at least one '|---' separator row.
    if "|---" in lowered or "| ---" in lowered:
        return False
    return True


def _convert_directory(target_dir: str, overwrite: bool) -> None:
    if not os.path.isdir(target_dir):
        print(f"[ERROR] Not a directory: {target_dir}")
        sys.exit(2)

    json_files = sorted(
        os.path.join(target_dir, f)
        for f in os.listdir(target_dir)
        if f.lower().endswith(".json")
    )
    if not json_files:
        print(f"[WARN] No .json files found in {target_dir!r}.")
        return

    converted = 0
    skipped_table = 0
    skipped_exists = 0
    failed = 0
    for jp in json_files:
        out_docx = os.path.splitext(jp)[0] + ".docx"
        if os.path.exists(out_docx) and not overwrite:
            print(f"[SKIP] {os.path.basename(out_docx)} already exists.")
            skipped_exists += 1
            continue

        try:
            with open(jp, "r", encoding="utf-8") as f:
                payload = json.load(f)
        except Exception as exc:
            print(f"[FAIL] {os.path.basename(jp)}: cannot parse JSON ({exc}).")
            failed += 1
            continue

        if not _is_text_page(payload):
            print(f"[SKIP] {os.path.basename(jp)}: looks like a table page "
                  f"(use json_to_word.py instead).")
            skipped_table += 1
            continue

        try:
            convert_json_to_docx(jp, out_docx)
            converted += 1
            print(f"[ OK ] {os.path.basename(jp)} -> "
                  f"{os.path.basename(out_docx)}")
        except Exception as exc:
            failed += 1
            print(f"[FAIL] {os.path.basename(jp)}: {exc}")

    print()
    print(f"Converted: {converted}")
    print(f"Skipped (table page):    {skipped_table}")
    print(f"Skipped (already exists): {skipped_exists}")
    print(f"Failed:    {failed}")


def main(argv: Optional[List[str]] = None) -> int:
    parser = argparse.ArgumentParser(
        description="Convert PP-StructureV3 text-heavy JSON to .docx with "
                    "an outer frame matching the genealogy table style.")
    parser.add_argument("input", nargs="?",
                        help="Input .json file (single-file mode).")
    parser.add_argument("output", nargs="?",
                        help="Output .docx file (defaults to <input>.docx).")
    parser.add_argument("-d", "--dir",
                        help="Process every text-heavy *.json in this "
                             "directory; outputs <name>.docx next to each.")
    parser.add_argument("--overwrite", action="store_true",
                        help="Overwrite existing .docx files in -d mode.")
    args = parser.parse_args(argv)

    if args.dir:
        _convert_directory(args.dir, overwrite=args.overwrite)
        return 0

    if not args.input:
        parser.print_help()
        return 1

    if not os.path.isfile(args.input):
        print(f"[ERROR] Input file not found: {args.input}")
        return 2

    out_docx = args.output or (os.path.splitext(args.input)[0] + ".docx")
    convert_json_to_docx(args.input, out_docx)
    print(f"Saved: {out_docx}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
