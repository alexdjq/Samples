# -*- coding: utf-8 -*-
"""
Convert PaddlePaddle PP-StructureV3 JSON output into a Word (.docx) file.

The script reads the JSON produced by PP-StructureV3 (which contains layout
information, table cell bounding boxes and OCR text boxes) and reconstructs
the table(s) into a Word document by mapping each recognized text box to the
table cell whose bounding box contains it.

Usage:
    python json_to_word.py <input_json> [output_docx]

Dependencies:
    pip install python-docx
"""

from __future__ import annotations

import json
import os
import sys
from typing import List, Tuple, Dict, Any, Optional

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Missing dependency 'python-docx'. Please install it via:")
    print("    pip install python-docx")
    sys.exit(1)


# ---------------------------------------------------------------------------
# Page / table layout constants  (A4 page, fixed 5 rows x 24 columns)
# ---------------------------------------------------------------------------

A4_WIDTH_CM  = 21.0
A4_HEIGHT_CM = 29.7

MARGIN_TOP_CM    = 0.5
MARGIN_BOTTOM_CM = 0.5
MARGIN_LEFT_CM   = 1.0
MARGIN_RIGHT_CM  = 1.0

TABLE_ROWS = 5
TABLE_COLS = 24

_USABLE_W_CM = A4_WIDTH_CM  - MARGIN_LEFT_CM  - MARGIN_RIGHT_CM   # 19 cm
_USABLE_H_CM = A4_HEIGHT_CM - MARGIN_TOP_CM   - MARGIN_BOTTOM_CM  # 27.7 cm

# Reserve a tiny safety buffer (a few mm) below the table so Word
# never spills the last row to a second page because of paragraph
# baseline / line-spacing rounding.
_TABLE_SAFETY_CM = 0.4
_TABLE_USABLE_H_CM = _USABLE_H_CM - _TABLE_SAFETY_CM     # 26.9 cm

CELL_W_CM = _USABLE_W_CM       / TABLE_COLS   # ~0.792 cm  (one Chinese char wide)
CELL_H_CM = _TABLE_USABLE_H_CM / TABLE_ROWS   # ~5.38 cm

FONT_SIZE_MAX_PT = 12.0
FONT_SIZE_MIN_PT = 7.0
FONT_SIZE_STEP   = 0.5


# ---------------------------------------------------------------------------
# Geometry helpers
# ---------------------------------------------------------------------------

Box = Tuple[float, float, float, float]  # (x1, y1, x2, y2)


def box_center(box: Box) -> Tuple[float, float]:
    x1, y1, x2, y2 = box
    return (x1 + x2) / 2.0, (y1 + y2) / 2.0


def point_in_box(px: float, py: float, box: Box, tol: float = 2.0) -> bool:
    x1, y1, x2, y2 = box
    return (x1 - tol) <= px <= (x2 + tol) and (y1 - tol) <= py <= (y2 + tol)


def box_overlap_area(a: Box, b: Box) -> float:
    ax1, ay1, ax2, ay2 = a
    bx1, by1, bx2, by2 = b
    inter_w = max(0.0, min(ax2, bx2) - max(ax1, bx1))
    inter_h = max(0.0, min(ay2, by2) - max(ay1, by1))
    return inter_w * inter_h


# ---------------------------------------------------------------------------
# Cell grid reconstruction
# ---------------------------------------------------------------------------

def build_grid_from_cells(cell_boxes: List[Box]) -> List[List[Box]]:
    """Group cell boxes into a 2D grid.

    Cluster cells by Y center to obtain rows, sort cells inside each row
    by X center to obtain columns, then return ``grid[row][col] -> box``.
    """
    if not cell_boxes:
        return []

    centers = [box_center(b) for b in cell_boxes]
    heights = [b[3] - b[1] for b in cell_boxes]
    avg_h = sum(heights) / len(heights)
    row_tol = max(avg_h * 0.4, 5.0)

    indexed = sorted(range(len(cell_boxes)), key=lambda i: centers[i][1])
    rows: List[List[int]] = []
    for i in indexed:
        cy = centers[i][1]
        placed = False
        for row in rows:
            row_cy = sum(centers[j][1] for j in row) / len(row)
            if abs(cy - row_cy) <= row_tol:
                row.append(i)
                placed = True
                break
        if not placed:
            rows.append([i])

    rows.sort(key=lambda r: sum(centers[j][1] for j in r) / len(r))
    for row in rows:
        row.sort(key=lambda j: centers[j][0])

    grid: List[List[Box]] = [[cell_boxes[j] for j in row] for row in rows]
    return grid


def assign_text_to_cells(
    grid: List[List[Box]],
    text_boxes: List[Box],
    texts: List[str],
) -> List[List[List[int]]]:
    """Assign indices of text fragments to the matching cell of ``grid``.

    Returns a 2D list ``bucket[row][col] -> [frag_idx, ...]``.
    """
    rows = len(grid)
    cols = max((len(r) for r in grid), default=0)
    bucket: List[List[List[int]]] = [
        [[] for _ in range(cols)] for _ in range(rows)
    ]

    for idx, (tb, txt) in enumerate(zip(text_boxes, texts)):
        if not txt or not txt.strip():
            continue
        cx, cy = box_center(tb)

        best = (-1, -1)
        best_score = -1.0
        for r_idx, row in enumerate(grid):
            for c_idx, cell in enumerate(row):
                if point_in_box(cx, cy, cell):
                    score = box_overlap_area(tb, cell) + 1e6
                else:
                    score = box_overlap_area(tb, cell)
                if score > best_score:
                    best_score = score
                    best = (r_idx, c_idx)

        if best_score <= 0:
            min_dist = float("inf")
            for r_idx, row in enumerate(grid):
                for c_idx, cell in enumerate(row):
                    ccx, ccy = box_center(cell)
                    d = (ccx - cx) ** 2 + (ccy - cy) ** 2
                    if d < min_dist:
                        min_dist = d
                        best = (r_idx, c_idx)

        r_idx, c_idx = best
        if 0 <= r_idx < rows and 0 <= c_idx < cols:
            bucket[r_idx][c_idx].append(idx)

    return bucket


# ---------------------------------------------------------------------------
# Vertical-Chinese cell content splitting
# ---------------------------------------------------------------------------

def split_into_vertical_lines(
    frag_indices: List[int],
    text_boxes: List[Box],
    texts: List[str],
) -> List[str]:
    """Split fragments inside a wide cell into reading lines.

    The PP-Structure OCR returns each piece of vertically-written Chinese
    text as a tall narrow box (height >> width).  Lines (i.e. "columns of
    glyphs" as printed in the original document) are then identified by
    grouping fragments whose horizontal centers are close enough.  Within
    one line, fragments are concatenated top-to-bottom.  Lines themselves
    are emitted right-to-left, which is the natural reading order for
    traditional vertical Chinese text.
    """
    if not frag_indices:
        return []

    items: List[Tuple[float, float, float, float, str]] = []
    widths: List[float] = []
    for idx in frag_indices:
        x1, y1, x2, y2 = text_boxes[idx]
        items.append((x1, y1, x2, y2, texts[idx]))
        widths.append(x2 - x1)

    if not items:
        return []

    widths.sort()
    median_w = widths[len(widths) // 2]
    col_tol = max(median_w * 0.7, 8.0)

    # Sort by x descending (right-to-left).
    items.sort(key=lambda t: -((t[0] + t[2]) / 2.0))

    columns: List[List[Tuple[float, float, float, float, str]]] = []
    for it in items:
        cx = (it[0] + it[2]) / 2.0
        placed = False
        for col in columns:
            col_cx = sum((c[0] + c[2]) / 2.0 for c in col) / len(col)
            if abs(cx - col_cx) <= col_tol:
                col.append(it)
                placed = True
                break
        if not placed:
            columns.append([it])

    # Right-to-left ordering of columns.
    columns.sort(key=lambda col: -sum((c[0] + c[2]) / 2.0 for c in col) / len(col))

    lines: List[str] = []
    for col in columns:
        col.sort(key=lambda c: c[1])  # top-to-bottom
        lines.append("".join(c[4] for c in col))
    return lines


# ---------------------------------------------------------------------------
# Word document helpers
# ---------------------------------------------------------------------------

def _set_run_chinese_font(run, font_name: str = "SimSun") -> None:
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)


def _clear_cell_text_direction(cell) -> None:
    """Remove any text-direction setting on this cell.

    The default left-to-right, top-to-bottom (lrTb) direction keeps all
    characters upright, which is what we want when laying out Chinese
    "vertical text" by putting one character per paragraph.  We never
    set tbRl here because that would cause Word to rotate Chinese
    glyphs in some font configurations.
    """
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        return
    td = tcPr.find(qn("w:textDirection"))
    if td is not None:
        tcPr.remove(td)


def _count_glyphs(text: str) -> Tuple[int, int]:
    """Return (max_segments_in_a_column, number_of_columns) for ``text``.

    Each ``\n``-separated piece is treated as one column, and within a
    column every character (Chinese, digit, etc.) counts as exactly one
    segment because the cell is laid out as one paragraph per character.
    """
    if not text:
        return 0, 0
    columns = [c for c in text.split("\n") if c]
    if not columns:
        return 0, 0
    max_segs = max(len(col) for col in columns)
    return max_segs, len(columns)


def _choose_font_size(text: str,
                      cell_w_cm: float = CELL_W_CM,
                      cell_h_cm: float = CELL_H_CM) -> float:
    """Pick the largest font (pt) that lets ``text`` fit one cell.

    Sizing model — must match :func:`set_cell_text`:
      * The cell uses lrTb text direction.
      * Each Chinese character occupies one paragraph (~ ``font_pt *
        0.0353`` cm tall and ~one em wide).
      * Each run of ASCII characters (a year, a range, ...) also
        occupies one paragraph but is rendered with a smaller, fitted
        font so that it always fits inside the cell width.  It still
        consumes one paragraph of vertical space, sized like a
        Chinese glyph.
      * Adjacent visual columns inside the same cell are separated by
        one empty paragraph.

    We therefore only need to make sure the *total number of
    paragraphs* fits inside the cell height, and that the cell width
    can hold one Chinese glyph.
    """
    if not text:
        return FONT_SIZE_MAX_PT

    columns = [c for c in text.split("\n") if c]
    if not columns:
        return FONT_SIZE_MAX_PT

    column_tokens = [_tokenize_for_vertical(c) for c in columns]
    column_tokens = [toks for toks in column_tokens if toks]
    if not column_tokens:
        return FONT_SIZE_MAX_PT

    total_paras = (
        sum(len(toks) for toks in column_tokens)
        + max(0, len(column_tokens) - 1)
    )

    inner_w = max(cell_w_cm - 0.05, 0.1)
    inner_h = max(cell_h_cm - 0.05, 0.1)

    size = FONT_SIZE_MAX_PT
    while size > FONT_SIZE_MIN_PT:
        em_cm = size * 0.0353
        height_need = total_paras * em_cm
        width_need  = em_cm  # one CJK glyph wide
        if height_need <= inner_h and width_need <= inner_w:
            break
        size -= FONT_SIZE_STEP
    return max(size, FONT_SIZE_MIN_PT)


def _tokenize_for_vertical(text: str) -> List[str]:
    """Split ``text`` into vertical-layout tokens.

    Rules:
      * Each Chinese character (or punctuation written in fullwidth /
        CJK form) becomes its own token so that it stands alone on one
        paragraph (visual vertical layout, glyph upright).
      * Consecutive ASCII / halfwidth characters — digits, letters,
        the Latin parentheses ``( )``, hyphen ``-``, dot ``.``, etc.
        — are merged into a SINGLE token.  This token is rendered as
        one horizontal paragraph, matching the original document where
        a year like ``1990`` is written as four digits side-by-side
        between vertical Chinese characters.

    Whitespace is dropped because the cell is a tightly packed
    vertical run of glyphs.
    """
    tokens: List[str] = []
    buf: List[str] = []
    for ch in text:
        if ch.isspace():
            if buf:
                tokens.append("".join(buf))
                buf = []
            continue
        # ASCII range (digits, letters, '(', ')', '-', '.', ',' etc.)
        # is treated as horizontal text that stays glued together.
        if ord(ch) < 128:
            buf.append(ch)
        else:
            if buf:
                tokens.append("".join(buf))
                buf = []
            tokens.append(ch)
    if buf:
        tokens.append("".join(buf))
    return tokens


def set_cell_text(cell, text: str, bold: bool = False,
                  cell_w_cm: float = CELL_W_CM,
                  cell_h_cm: float = CELL_H_CM,
                  center: bool = False) -> None:
    """Render ``text`` inside a Word cell as upright vertical text.

    Layout rules (matching the reference document):
      * The cell uses the default (lrTb) text direction, so no glyph
        is rotated.  Chinese characters and digits all stay upright.
      * Each Chinese character is placed on its own paragraph, giving
        the visual effect of vertical Chinese.
      * Runs of ASCII characters (e.g. ``1990``, ``(1949-1995)``) are
        kept together on a single horizontal paragraph so that, just
        like in the reference document, the four digits of a year
        share one row instead of being split into four rows.  An
        ASCII run's font is shrunk so it fits in the single-glyph
        cell width.
      * ``\n`` in the source separates *visual columns* inside the
        same cell (e.g. multiple persons).  An empty paragraph marks
        the gap between adjacent columns.
      * Font size is auto-shrunk **per cell** so that long content
        only affects the offending cell, not its neighbours.
    """
    cell.text = ""
    # Make sure no left-over vertical (tbRl) text direction is applied
    # so that Chinese characters do NOT get rotated by Word.
    _clear_cell_text_direction(cell)
    font_size = _choose_font_size(text, cell_w_cm, cell_h_cm)

    columns = text.split("\n") if text else [""]
    first = True
    last_para = None  # last paragraph that received real content

    def _new_para():
        nonlocal first
        if first:
            first = False
            return cell.paragraphs[0]
        return cell.add_paragraph()

    def _emit_token(token: str) -> None:
        nonlocal last_para
        p = _new_para()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(token)
        # ASCII-only tokens (e.g. years like 1990 or ranges like
        # (1949-1995)) are rendered horizontally inside ONE paragraph.
        # We shrink their font so the whole run fits in the
        # single-glyph-wide cell, mimicking the reference document
        # where year digits are noticeably smaller than the
        # surrounding Chinese characters.
        if token and all(ord(c) < 128 for c in token):
            # Year tokens like ``1990`` and range tokens like
            # ``(1949-1995)`` are rendered horizontally on a single
            # line.  We previously shrank them aggressively (~45 % of
            # the base size for a 4-digit year) which made them hard
            # to read.  Use a gentler shrink so that, e.g., ``1990``
            # ends up at ~69 % of the surrounding Chinese characters
            # while still fitting inside the single-glyph cell width.
            ascii_size = font_size * min(1.0, 1.0 / (0.36 * len(token)))
            ascii_size = max(ascii_size, FONT_SIZE_MIN_PT)
            run.font.size = Pt(ascii_size)
        else:
            run.font.size = Pt(font_size)
        run.font.name = "SimSun"
        run.bold = bold
        _set_run_chinese_font(run, "SimSun")
        last_para = p

    # First pass: walk columns and convert *runs of empty columns* into
    # an explicit ``space_after`` (in points) that we attach to the
    # previous non-empty paragraph.  Using a real, measurable point
    # value is far more reliable than relying on empty paragraphs
    # whose height Word tends to collapse.
    pending_gap_units = 0  # number of empty columns waiting to be flushed

    def _flush_gap() -> None:
        """Apply the accumulated gap (in units of one CJK glyph
        height) to ``last_para``'s ``space_after``.  If there is no
        previous content paragraph yet, emit a single blank-height
        paragraph so the gap sits at the very top of the cell."""
        nonlocal pending_gap_units, last_para
        if pending_gap_units <= 0:
            return
        gap_pt = pending_gap_units * font_size
        if last_para is not None:
            existing = last_para.paragraph_format.space_after
            base = existing.pt if existing is not None else 0
            last_para.paragraph_format.space_after = Pt(base + gap_pt)
        else:
            # Gap appears before any real content: emit a spacer
            # paragraph with the requested height as space_before.
            p = _new_para()
            p.paragraph_format.space_before = Pt(gap_pt)
            p.paragraph_format.space_after  = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            if center:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("")
            run.font.size = Pt(font_size)
            run.font.name = "SimSun"
            _set_run_chinese_font(run, "SimSun")
        pending_gap_units = 0

    for col_idx, col_text in enumerate(columns):
        if col_text == "":
            # Each empty column contributes one CJK-glyph-height gap.
            pending_gap_units += 1
            continue

        # For two adjacent NON-empty columns (no empty columns between
        # them), keep the legacy 1-glyph separator gap so existing
        # multi-person cells still look the same.
        if (col_idx > 0
                and columns[col_idx - 1] != ""
                and pending_gap_units == 0):
            pending_gap_units = 1

        _flush_gap()

        for token in _tokenize_for_vertical(col_text):
            _emit_token(token)

    # If the text ended with empty columns, flush the trailing gap so
    # it still produces visible whitespace at the bottom of the cell.
    _flush_gap()

    # When ``center`` is requested, vertically center the whole
    # paragraph stack inside the tall cell (used by the rightmost
    # generation-label column to make 十六世/十七世 sit in the
    # middle of their cell, matching the original document).
    cell.vertical_alignment = (
        WD_ALIGN_VERTICAL.CENTER if center else WD_ALIGN_VERTICAL.TOP
    )


# ---------------------------------------------------------------------------
# Page and fixed-table helpers
# ---------------------------------------------------------------------------

def _setup_a4_page(doc: Document) -> None:
    """Configure the document for A4 with tight, equal margins."""
    section = doc.sections[0]
    section.page_width    = Cm(A4_WIDTH_CM)
    section.page_height   = Cm(A4_HEIGHT_CM)
    section.top_margin    = Cm(MARGIN_TOP_CM)
    section.bottom_margin = Cm(MARGIN_BOTTOM_CM)
    section.left_margin   = Cm(MARGIN_LEFT_CM)
    section.right_margin  = Cm(MARGIN_RIGHT_CM)


def _set_table_fixed_layout(table) -> None:
    """Force fixed column widths so Word does not auto-resize."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    layout = tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")


def _zero_table_cell_margins(table) -> None:
    """Set all default cell margins (top/bottom/left/right) to 0.

    By default Word adds ~0.05cm of inner padding on every side of
    every cell, which adds up across 5 rows and pushes the table to a
    second page.  Zeroing them out keeps the whole table on one page.
    """
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    cell_mar = tblPr.find(qn("w:tblCellMar"))
    if cell_mar is None:
        cell_mar = OxmlElement("w:tblCellMar")
        tblPr.append(cell_mar)
    for side in ("top", "left", "bottom", "right"):
        node = cell_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            cell_mar.append(node)
        node.set(qn("w:w"),    "0")
        node.set(qn("w:type"), "dxa")


def _set_row_exact_height(row, height_cm: float) -> None:
    """Force an exact row height (hRule=exact, no auto-grow)."""
    tr = row._tr
    trPr = tr.find(qn("w:trPr"))
    if trPr is None:
        trPr = OxmlElement("w:trPr")
        tr.insert(0, trPr)
    trH = trPr.find(qn("w:trHeight"))
    if trH is None:
        trH = OxmlElement("w:trHeight")
        trPr.append(trH)
    trH.set(qn("w:val"),   str(int(height_cm * 567)))  # 1 cm = 567 twips
    trH.set(qn("w:hRule"), "exact")


def _set_cell_width(cell, width_cm: float) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"),    str(int(width_cm * 567)))
    tcW.set(qn("w:type"), "dxa")


def _collapse_empty_columns(table, min_visible_cm: float = 0.05) -> None:
    """Shrink every column whose every cell is empty to ``min_visible_cm``.

    Why this matters: under the global X-anchor strategy used in
    :func:`render_table`, dense rows (e.g. 十世 with 12 persons) and
    sparse rows (e.g. 九世 with 6 persons) share the same column grid.
    The result is large stretches of completely-empty columns between
    sparse-row entries, which read as a "huge gap" to the user.

    By detecting columns where *every* row is blank and physically
    shrinking them to a near-zero width, we keep the cross-row anchor
    alignment intact (the order and relative position of person
    columns is preserved) while removing the visual dead-space.  The
    width freed by collapsing those columns is redistributed evenly
    across the columns that actually carry content, so the table still
    fills the full :data:`_USABLE_W_CM`.

    Parameters
    ----------
    table : docx.table.Table
        The table to compact.  Must be built with a fixed layout
        (`_set_table_fixed_layout`) for the new widths to take effect.
    min_visible_cm : float, optional
        Width to assign to each empty column.  Defaults to 0.05 cm,
        small enough to be invisible but non-zero so Word does not
        drop the column entirely.
    """
    n_rows = len(table.rows)
    n_cols = len(table.columns)
    if n_rows == 0 or n_cols == 0:
        return

    # 1) Detect empty columns.
    empty_flags: List[bool] = []
    for c in range(n_cols):
        is_empty = True
        for r in range(n_rows):
            if table.cell(r, c).text.strip():
                is_empty = False
                break
        empty_flags.append(is_empty)

    n_empty = sum(empty_flags)
    if n_empty == 0 or n_empty == n_cols:
        # Nothing to compact (no empties, or every column empty).
        return

    n_filled = n_cols - n_empty
    width_for_empty = min_visible_cm
    width_for_filled = (
        _USABLE_W_CM - width_for_empty * n_empty
    ) / n_filled
    # Defensive lower bound: never let filled columns shrink below the
    # empty-column width.
    if width_for_filled < min_visible_cm:
        width_for_filled = min_visible_cm

    # 2) Update every cell's tcW so Word renders the new widths.
    for r in range(n_rows):
        for c in range(n_cols):
            new_w = width_for_empty if empty_flags[c] else width_for_filled
            _set_cell_width(table.cell(r, c), new_w)

    # 3) Update <w:tblGrid> too.  Under fixed-layout tables Word also
    #    consults the table grid; keeping it in sync with tcW prevents
    #    some renderers from falling back to equal column widths.
    tbl = table._tbl
    tblGrid = tbl.find(qn("w:tblGrid"))
    if tblGrid is not None:
        # Remove existing gridCol entries, then re-create them so the
        # count always matches ``n_cols`` exactly.
        for gc in list(tblGrid.findall(qn("w:gridCol"))):
            tblGrid.remove(gc)
        for c in range(n_cols):
            new_w = width_for_empty if empty_flags[c] else width_for_filled
            gc = OxmlElement("w:gridCol")
            gc.set(qn("w:w"), str(int(new_w * 567)))
            tblGrid.append(gc)


def _set_cell_borders(cell,
                      top: str = "single", bottom: str = "single",
                      left: str = "nil",   right: str = "nil",
                      sz: str = "4") -> None:
    """Set per-cell borders.

    Each side accepts a Word border-style token, e.g. ``"single"`` for a
    visible line or ``"nil"`` to hide that side.  ``sz`` is the line
    weight in eighths of a point (4 == 0.5pt, the Word default).
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for side, val in (("top", top), ("left", left),
                      ("bottom", bottom), ("right", right)):
        node = tcBorders.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tcBorders.append(node)
        node.set(qn("w:val"),   val)
        node.set(qn("w:sz"),    sz)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "000000")


def _apply_table_border_style(table, header_col: int,
                              merged_col: int = -1) -> None:
    """Apply a "key-borders-only" style typical of Chinese genealogies.

    The original PDF for this project shows three structural lines:

    1. The full outer rectangle around the whole table.
    2. A vertical line on the LEFT side of ``header_col``, which
       separates the right-hand "page-edge band" (generation labels
       and the merged title column) from the body of person columns.
    3. Horizontal separators between successive 世代 rows so each
       generation reads as its own band.

    All other inside-vertical lines (between adjacent person columns)
    are suppressed -- those are precisely the lines that turned the
    earlier output into an unwanted grid.

    Parameters
    ----------
    table : docx.table.Table
        Table to style.
    header_col : int
        Index of the column whose **left** side gets a vertical
        separator drawn.  Pass a negative value or 0 to skip drawing
        the separator (when there is no header band at all).
    merged_col : int, optional
        Index of a vertically-merged column (e.g. the right-edge
        title column ``"渑池/杜氏宗谱/page-number"``).  Horizontal
        separators between rows are *not* drawn over this column so
        the merged cell reads as one continuous vertical band, just
        like in the PDF.  Pass ``-1`` (default) when no column is
        merged.
    """
    n_rows = len(table.rows)
    n_cols = len(table.columns)
    if n_rows == 0 or n_cols == 0:
        return

    # ------------------------------------------------------------------
    # Step 1: hide every cell border, then re-enable only the ones we
    # actually want.  Working at the cell level (rather than the
    # table-level <w:tblBorders>) gives us per-cell control which is
    # required for the "title column separator" and "between-row
    # horizontal lines" rules.
    # ------------------------------------------------------------------
    for r in range(n_rows):
        for c in range(n_cols):
            _set_cell_borders(table.cell(r, c),
                              top="nil", bottom="nil",
                              left="nil", right="nil")

    # Suppress the table-level inside borders too, otherwise some Word
    # versions still draw the implicit "Table Grid" lines.
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = tblPr.find(qn("w:tblBorders"))
    if tblBorders is None:
        tblBorders = OxmlElement("w:tblBorders")
        tblPr.append(tblBorders)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = tblBorders.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tblBorders.append(node)
        node.set(qn("w:val"),   "nil")
        node.set(qn("w:sz"),    "0")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "auto")

    # ------------------------------------------------------------------
    # Step 2: draw the outer rectangle.
    # ------------------------------------------------------------------
    last_row = n_rows - 1
    last_col = n_cols - 1
    for c in range(n_cols):
        # Top edge
        top_cell = table.cell(0, c)
        _set_cell_borders(top_cell,
                          top="single",
                          bottom=_read_cell_border(top_cell, "bottom"),
                          left=_read_cell_border(top_cell, "left"),
                          right=_read_cell_border(top_cell, "right"))
        # Bottom edge
        bot_cell = table.cell(last_row, c)
        _set_cell_borders(bot_cell,
                          top=_read_cell_border(bot_cell, "top"),
                          bottom="single",
                          left=_read_cell_border(bot_cell, "left"),
                          right=_read_cell_border(bot_cell, "right"))
    for r in range(n_rows):
        # Left edge
        left_cell = table.cell(r, 0)
        _set_cell_borders(left_cell,
                          top=_read_cell_border(left_cell, "top"),
                          bottom=_read_cell_border(left_cell, "bottom"),
                          left="single",
                          right=_read_cell_border(left_cell, "right"))
        # Right edge
        right_cell = table.cell(r, last_col)
        _set_cell_borders(right_cell,
                          top=_read_cell_border(right_cell, "top"),
                          bottom=_read_cell_border(right_cell, "bottom"),
                          left=_read_cell_border(right_cell, "left"),
                          right="single")

    # ------------------------------------------------------------------
    # Step 3: vertical separator on the LEFT side of ``header_col``.
    # Skip when header_col <= 0 (would coincide with the outer left
    # edge already drawn) or >= n_cols (out of range).
    # ------------------------------------------------------------------
    if 0 < header_col < n_cols:
        for r in range(n_rows):
            cell = table.cell(r, header_col)
            _set_cell_borders(cell,
                              top=_read_cell_border(cell, "top"),
                              bottom=_read_cell_border(cell, "bottom"),
                              left="single",
                              right=_read_cell_border(cell, "right"))

    # ------------------------------------------------------------------
    # Step 4: horizontal lines between successive generation rows.
    # We draw them as the BOTTOM border of every row except the last
    # (whose bottom is already covered by the outer rectangle).  We
    # deliberately draw across **every** column, including the merged
    # title column on the far right -- the original PDF shows the
    # row separator running edge-to-edge, cutting through the title
    # band as well.  ``merged_col`` is currently unused but kept in
    # the signature for forward compatibility.
    # ------------------------------------------------------------------
    _ = merged_col  # noqa: F841 -- accepted but intentionally unused.
    for r in range(n_rows - 1):
        for c in range(n_cols):
            cell = table.cell(r, c)
            _set_cell_borders(cell,
                              top=_read_cell_border(cell, "top"),
                              bottom="single",
                              left=_read_cell_border(cell, "left"),
                              right=_read_cell_border(cell, "right"))


def _read_cell_border(cell, side: str) -> str:
    """Return the current border value of ``side`` ('nil' if unset)."""
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        return "nil"
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        return "nil"
    node = tcBorders.find(qn(f"w:{side}"))
    if node is None:
        return "nil"
    return node.get(qn("w:val")) or "nil"


def _create_fixed_table(doc: Document, n_cols: int = TABLE_COLS):
    """Create a TABLE_ROWS x ``n_cols`` table.

    ``n_cols`` defaults to the canonical :data:`TABLE_COLS` value but
    callers can request a wider table when a single page contains more
    person-columns than the default 24-column layout can hold.  The
    column width is scaled so the whole table still fits within the
    A4 usable width (no horizontal overflow).
    """
    if n_cols < 1:
        n_cols = TABLE_COLS
    cell_w_cm = _USABLE_W_CM / n_cols
    table = doc.add_table(rows=TABLE_ROWS, cols=n_cols)
    # NOTE: deliberately do *not* set ``table.style = "Table Grid"``.
    # The Grid style applies a default 0.5pt border on every side of
    # every cell which some renderers (notably WPS / older Word
    # versions) still draw even when ``tblBorders`` are set to ``nil``
    # at the table level.  Using the default ("Normal Table") style and
    # then explicitly hiding every border in :func:`_apply_table_border_style`
    # gives the most reliable borderless rendering across viewers.
    _set_table_fixed_layout(table)
    _zero_table_cell_margins(table)
    for r_idx in range(TABLE_ROWS):
        _set_row_exact_height(table.rows[r_idx], CELL_H_CM)
        for c_idx in range(n_cols):
            _set_cell_width(table.cell(r_idx, c_idx), cell_w_cm)
    # Default: rightmost column is the generation-label column.
    _apply_table_border_style(table, header_col=n_cols - 1)
    return table


# ---------------------------------------------------------------------------
# Table rendering
# ---------------------------------------------------------------------------

def _split_cell_into_person_lines(
    frag_indices: List[int],
    t_boxes: List[Box],
    t_texts: List[str],
) -> List[Tuple[float, str]]:
    """Split a wide cell into "person columns".

    Returns a list of ``(x_center, text)`` tuples, one per detected
    vertical text column inside the cell.  The text is the top-to-bottom
    concatenation of all fragments that belong to the same X cluster.
    The list is **not** pre-sorted; the caller decides the final order.
    """
    if not frag_indices:
        return []

    items: List[Tuple[float, float, float, float, str]] = []
    widths: List[float] = []
    for idx in frag_indices:
        x1, y1, x2, y2 = t_boxes[idx]
        items.append((x1, y1, x2, y2, t_texts[idx]))
        widths.append(x2 - x1)

    widths.sort()
    median_w = widths[len(widths) // 2]
    col_tol = max(median_w * 0.7, 8.0)

    # Group fragments by X center.
    items.sort(key=lambda t: (t[0] + t[2]) / 2.0)
    groups: List[List[Tuple[float, float, float, float, str]]] = []
    for it in items:
        cx = (it[0] + it[2]) / 2.0
        placed = False
        for g in groups:
            g_cx = sum((c[0] + c[2]) / 2.0 for c in g) / len(g)
            if abs(cx - g_cx) <= col_tol:
                g.append(it)
                placed = True
                break
        if not placed:
            groups.append([it])

    out: List[Tuple[float, str]] = []
    for g in groups:
        g.sort(key=lambda c: c[1])  # top-to-bottom
        text = "".join(c[4] for c in g)
        x_center = sum((c[0] + c[2]) / 2.0 for c in g) / len(g)
        out.append((x_center, text))
    return out


def _cluster_x_anchors(
    x_values: List[float], tol: float
) -> List[float]:
    """Cluster x values into anchor points (sorted descending)."""
    if not x_values:
        return []
    sorted_x = sorted(x_values)
    clusters: List[List[float]] = [[sorted_x[0]]]
    for x in sorted_x[1:]:
        if abs(x - clusters[-1][-1]) <= tol:
            clusters[-1].append(x)
        else:
            clusters.append([x])
    anchors = [sum(c) / len(c) for c in clusters]
    anchors.sort(reverse=True)  # right -> left
    return anchors


# ---------------------------------------------------------------------------
# Page title / page-number extraction
# ---------------------------------------------------------------------------

_CN_DIGIT_CHARS = set("〇○零一二三四五六七八九十百千两")
# Characters that strongly identify a multi-digit page number (e.g. the
# 0 in "三〇八").  A bare generation index such as "十八" never contains
# any of these, so requiring at least one of them lets us tell page
# numbers apart from generation labels that lost their trailing "世".
_CN_PAGE_REQUIRED_CHARS = set("〇○零百千")


def _is_cn_page_number(text: str) -> bool:
    """Return True if ``text`` looks like a Chinese-numeral page number.

    Page numbers in this document are written vertically as Chinese
    digits, e.g. "三〇八".  To avoid misclassifying a generation
    index such as "十八" (a truncated "十八世") we require:
      * length 2-5, all CJK digit / counter glyphs;
      * AND at least one "page-only" glyph (〇 / ○ / 零 / 百 / 千).
    """
    if not text:
        return False
    t = text.strip()
    if not (2 <= len(t) <= 5):
        return False
    if not all(ch in _CN_DIGIT_CHARS for ch in t):
        return False
    return any(ch in _CN_PAGE_REQUIRED_CHARS for ch in t)


def _extract_title_and_page_for_page(
    page: Dict[str, Any]
) -> Tuple[str, str]:
    """Extract "渑池" / "杜氏宗谱" title fragments and the page number
    for **one** page only.

    These three pieces of text live on the page edge of the source PDF
    (vertical strip on the right hand side of the genealogy table) and
    are NOT part of the table cell grid, so they would otherwise be
    dropped.  We pull them out of this page's ``overall_ocr_res.rec_texts``
    so that every page renders its own page number ("三〇六", "三〇七",
    "三〇八", "三〇九" ...), instead of all pages sharing the first one
    found in the document.

    Returns ``(title_text, page_text)`` where ``title_text`` is the
    non-empty title fragments joined by ``\n`` (so the caller can render
    each as its own person-column inside the merged side strip cell)
    and ``page_text`` is the page number string (or an empty string if
    not detected on this page).
    """
    pruned = page.get("prunedResult", page)
    overall = pruned.get("overall_ocr_res", {})

    title_parts: List[str] = []
    page_text = ""
    seen: set = set()
    for txt in overall.get("rec_texts", []):
        if not txt:
            continue
        t = txt.strip()
        if not t or t in seen:
            continue
        # Title fragments — match the originals ("渑池", "杜氏宗谱").
        if t in ("渑池", "杜氏宗谱"):
            title_parts.append(t)
            seen.add(t)
            continue
        # Page number: take the first CN-digit-only fragment found on
        # this page.  Skip strings that merely contain "渑池" mid-word
        # (e.g. "保玲1980年生适渑池") - those are not page numbers.
        if not page_text and _is_cn_page_number(t):
            page_text = t
            seen.add(t)
    # Preserve the natural "渑池" before "杜氏宗谱" order if both were found.
    title_parts.sort(key=lambda s: 0 if s == "渑池" else 1)
    return "\n".join(title_parts), page_text


def _filter_main_table_cells(cell_box_list: List[Box]) -> List[Box]:
    """Drop OCR cells that belong to the page-edge side strip rather
    than the genealogy table itself.

    Some pages (e.g. the first sheet of 样章.pdf) have their right-
    hand "渑池 / 杜氏宗谱 / 三〇六" vertical strip mis-detected as part
    of the table; that injects a few abnormally tall narrow cells that
    span almost the full page height (height >> the regular row height).
    Those cells:

      * break the row clustering inside :func:`build_grid_from_cells`
        because their Y center lies in the middle of multiple rows;
      * are NOT part of the 5-row genealogy grid and therefore must
        not be rendered as table content.

    We identify them by looking at the **median row height** of the
    grid.  Any cell whose height is significantly larger (≥ 2× the
    median) is treated as a page-edge artefact and removed.  Cells
    that lie horizontally outside the X-range of the regular rows are
    also dropped.
    """
    if not cell_box_list:
        return []

    heights = sorted(b[3] - b[1] for b in cell_box_list)
    median_h = heights[len(heights) // 2]
    if median_h <= 0:
        return list(cell_box_list)

    # Page-edge strip cells are dramatically taller than a normal row.
    kept = [b for b in cell_box_list if (b[3] - b[1]) < median_h * 2.0]
    if len(kept) < 2:
        # Filtering would erase too much - keep the original list and
        # let the downstream renderer try its best.
        return list(cell_box_list)

    # If after filtering we have a clean 5-row x 2-col layout (i.e. an
    # even count divisible by 5 with the narrow + wide column pair we
    # expect), drop any remaining outliers that sit outside the X
    # range covered by the bulk of the cells (page-edge strip cells
    # that happened to have the same height).
    xs = sorted([b[0] for b in kept] + [b[2] for b in kept])
    x_lo = xs[len(xs) // 20] if len(xs) >= 20 else xs[0]
    x_hi = xs[-len(xs) // 20 - 1] if len(xs) >= 20 else xs[-1]
    bulk = [b for b in kept if b[0] >= x_lo - 5 and b[2] <= x_hi + 5]
    if len(bulk) < 2:
        bulk = kept

    # ------------------------------------------------------------------
    # Drop **row-spanning merged cells**.
    #
    # On some pages (e.g. page 2 of 窑湾.pdf) the leftmost column is
    # detected as TWO tall cells that each vertically span 2 of the 5
    # generation rows.  Their height (~516 / ~709) is below the
    # ``2 * median`` threshold above, so they survive the first pass --
    # but they still ruin the row clustering inside
    # :func:`build_grid_from_cells`, splitting the genealogy grid into
    # 6 pseudo-rows instead of 5.  The very last real row (e.g. 十世)
    # is then pushed past ``TABLE_ROWS`` and silently dropped.
    #
    # We detect such cells by checking how much vertical overlap each
    # cell has with the "row body" of every other cell.  Any cell that
    # *contains* the Y center of two or more other cells is treated as
    # a merged super-cell and removed -- its contents will instead be
    # picked up via the global OCR text boxes when the renderer falls
    # back to its anchor-based layout.
    # ------------------------------------------------------------------
    def _y_center(b: Box) -> float:
        return (b[1] + b[3]) / 2.0

    centers_y = [_y_center(b) for b in bulk]
    pruned: List[Box] = []
    for i, b in enumerate(bulk):
        # A genuine *row-spanning* merged cell vertically swallows
        # cells from MULTIPLE other rows.  We detect this by counting
        # how many other cells have a Y center that is (a) inside this
        # cell's vertical span, AND (b) far enough from this cell's
        # own Y center to clearly belong to a different row -- the
        # threshold is half the cell's own height, which separates
        # same-row neighbours (almost identical Y centers) from cells
        # that genuinely sit in a different generation row.
        b_cy = _y_center(b)
        b_h = max(1.0, b[3] - b[1])
        spans = 0
        for j, cy in enumerate(centers_y):
            if j == i:
                continue
            if not (b[1] + 5 < cy < b[3] - 5):
                continue
            if abs(cy - b_cy) <= b_h * 0.25:
                # Cell ``j`` is on the same row as ``i`` (their Y
                # centers nearly coincide); ignore.
                continue
            spans += 1
        if spans >= 2:
            # This cell vertically swallows >=2 cells from other rows
            # -- it is a merged super-cell, drop it so the genealogy
            # grid keeps its clean 5-row structure.
            continue
        pruned.append(b)

    if len(pruned) >= 2:
        return pruned
    return bulk


def render_table(
    doc: Document,
    cell_box_list: List[Box],
    t_texts: List[str],
    t_boxes: List[Box],
    title_text: str = "",
    page_text: str = "",
) -> bool:
    """Render one detected table to ``doc``.

    Two layouts are supported:

    * **Genealogy (rotated) layout** — one narrow header column (e.g.
      "十六世~二十世") plus one wide content column per generation, with
      vertically written, right-to-left Chinese text.  The renderer:
        1. Treats each image-row as one **generation**.
        2. Splits the wide content cell into per-person vertical lines.
        3. Clusters those lines' X centers across **all** generations to
           form a global set of person-column anchors.
        4. Builds a Word table whose **rows are generations** ordered
           top-to-bottom (十六世 at the top), and whose **columns are
           aligned person columns** ordered left-to-right by image X.
           The generation label sits in the **rightmost** column,
           matching the original document layout.
    * **Generic grid** — falls back to a simple row x col rendering.
    """
    grid = build_grid_from_cells(cell_box_list)
    if not grid:
        return False

    bucket = assign_text_to_cells(grid, t_boxes, t_texts)

    rows = len(grid)
    cols = max(len(r) for r in grid)

    # ---- Detect "header column" layout (one column much narrower) -------
    transposed_layout = False
    header_col_idx: Optional[int] = None
    if cols >= 2:
        col_widths: List[float] = []
        for c in range(cols):
            ws = [grid[r][c][2] - grid[r][c][0] for r in range(rows) if c < len(grid[r])]
            col_widths.append(sum(ws) / len(ws) if ws else 0.0)
        max_w = max(col_widths)
        min_w = min(col_widths)
        if max_w > 0 and min_w / max_w < 0.25:
            header_col_idx = col_widths.index(min_w)
            transposed_layout = True

    if transposed_layout and header_col_idx is not None:
        # 1. Per generation: header label + list of (x_center, text).
        gen_count = rows
        header_labels: List[str] = []
        person_lists: List[List[Tuple[float, str]]] = []

        for r in range(gen_count):
            head_indices = bucket[r][header_col_idx] if header_col_idx < len(bucket[r]) else []
            head_lines = split_into_vertical_lines(head_indices, t_boxes, t_texts)
            header_labels.append("\n".join(head_lines).strip())

            persons: List[Tuple[float, str]] = []
            for c in range(cols):
                if c == header_col_idx:
                    continue
                if c < len(bucket[r]):
                    persons.extend(
                        _split_cell_into_person_lines(bucket[r][c], t_boxes, t_texts)
                    )
            person_lists.append(persons)

        # 2. Build global X anchors across all generations.
        all_x = [x for plist in person_lists for x, _ in plist]
        if not all_x:
            return False

        # Estimate tolerance from the *minimum non-zero* gap between
        # neighbouring person-column X centers across all generations.
        # The original document uses a fixed character pitch, so genuinely
        # different person-columns are separated by roughly one character
        # width.  Using half of that as the clustering tolerance makes the
        # algorithm robust to small OCR jitter while still keeping
        # neighbouring person columns distinct.
        widths_all = [t_boxes[i][2] - t_boxes[i][0] for i in range(len(t_boxes))]
        widths_all.sort()
        median_w = widths_all[len(widths_all) // 2] if widths_all else 30.0
        anchor_tol = max(median_w * 0.45, 10.0)
        anchors = _cluster_x_anchors(all_x, anchor_tol)  # right -> left

        # 3. Word table layout (transposed):
        #    - rows         = generations (one row per 世), ordered
        #                     top-to-bottom 十六世 -> 二十世 (i.e. by
        #                     image Y ascending).
        #    - cols 0..N-1  = person columns aligned to global anchors,
        #                     ordered left-to-right (= image X ascending).
        #                     This mirrors the original document where the
        #                     leftmost person column in the image becomes
        #                     the leftmost column in the Word table.
        #    - col   N      = generation label (header column on the
        #                     **right** side, matching the original
        #                     document layout).
        gen_order = sorted(
            range(gen_count),
            key=lambda r: (grid[r][0][1] + grid[r][0][3]) / 2.0,
        )

        # 4. For each generation, assign every (x, text) to the closest
        #    anchor index. ``anchors`` is sorted right-to-left (X
        #    descending), so to lay them out left-to-right in the Word
        #    table we map anchor index ``a_idx`` to Word column
        #    ``(n_anchor_cols - 1) - a_idx``.
        def _anchor_index(x: float) -> int:
            best_idx = 0
            best_d = float("inf")
            for i, a in enumerate(anchors):
                d = abs(x - a)
                if d < best_d:
                    best_d = d
                    best_idx = i
            return best_idx

        n_anchor_cols = len(anchors)
        word_rows = len(gen_order)
        word_cols = n_anchor_cols + 1  # anchor columns + trailing header
        header_col = n_anchor_cols      # rightmost column
        cell_text: List[List[str]] = [
            ["" for _ in range(word_cols)] for _ in range(word_rows)
        ]

        for w_row, g_idx in enumerate(gen_order):
            # Generation label goes into the rightmost column.
            cell_text[w_row][header_col] = header_labels[g_idx]

            # Sort persons by X ascending so that, when multiple persons
            # collapse onto the same anchor column, the leftmost one is
            # appended first (matching left-to-right reading).
            persons_sorted = sorted(person_lists[g_idx], key=lambda t: t[0])
            for x, text in persons_sorted:
                a_idx = _anchor_index(x)
                col = (n_anchor_cols - 1) - a_idx
                if cell_text[w_row][col]:
                    cell_text[w_row][col] += "\n" + text
                else:
                    cell_text[w_row][col] = text

        # 5. Render into a TABLE_ROWS x ``table_cols`` table.
        #
        #    Layout (right -> left):
        #      * column table_cols-1 (rightmost): book title "渑池 / 杜氏
        #        宗谱" on the first row + page number "三〇八" on the
        #        last row, mirroring the page-edge column of the
        #        original document.
        #      * column table_cols-2 : generation labels (十六世 ...).
        #      * columns 0 .. table_cols-3 : person columns, right-
        #        aligned to the generation column.
        #    When the title column is not requested (no extracted text),
        #    we fall back to the previous behaviour where the generation
        #    label sits in the rightmost column.
        #
        #    ``table_cols`` is normally the canonical ``TABLE_COLS``
        #    (24) but is automatically widened when the current page
        #    has so many distinct person-columns that the leftmost
        #    one(s) would otherwise be cropped off the table.  This
        #    keeps every OCR fragment visible in its correct image-
        #    relative position even when the OCR detected an unusually
        #    high number of vertical text columns on the page.
        has_title_col = bool(title_text or page_text)
        # Required column count = N person columns + 1 generation
        # label column + (1 title column if any).
        required_cols = word_cols + (1 if has_title_col else 0)
        table_cols = max(TABLE_COLS, required_cols)

        word_table = _create_fixed_table(doc, n_cols=table_cols)

        title_col_idx = table_cols - 1 if has_title_col else -1
        # Anchor for the generation-label column.
        gen_label_col = table_cols - 2 if has_title_col else table_cols - 1
        # Person columns occupy [0 .. gen_label_col-1]; right-align.
        right_col_offset = gen_label_col - (word_cols - 1)

        # Per-cell width depends on the (possibly widened) table.
        cell_w_cm_local = _USABLE_W_CM / table_cols

        # Re-apply borders to match the new header column position so
        # the only internal vertical line stays adjacent to the
        # generation-label column.  ``merged_col`` is the rightmost
        # title column (when present) which is vertically merged
        # across all rows; it must be excluded from the per-row
        # horizontal separators.
        _apply_table_border_style(
            word_table,
            header_col=gen_label_col,
            merged_col=title_col_idx,
        )

        # If the title column is enabled, vertically merge all 5 rows
        # of the rightmost column into a single tall cell.  The merged
        # cell holds "渑池 / 杜氏宗谱" near the top and the page number
        # near the bottom, with both vertically centered together so
        # they read as a single page-edge label, just like in the
        # source document.
        if has_title_col:
            top_cell = word_table.cell(0, title_col_idx)
            for r_idx in range(1, TABLE_ROWS):
                top_cell.merge(word_table.cell(r_idx, title_col_idx))

        for r_idx in range(TABLE_ROWS):
            for c_idx in range(table_cols):
                if c_idx == title_col_idx:
                    # The merged title/page cell is filled exactly
                    # once below; skip the per-row iteration here.
                    continue

                txt = ""
                bold_flag = False
                center_flag = False

                src_c = c_idx - right_col_offset
                if 0 <= r_idx < word_rows and 0 <= src_c < word_cols:
                    txt = cell_text[r_idx][src_c]
                    bold_flag = (src_c == header_col)
                if c_idx == gen_label_col:
                    center_flag = True

                set_cell_text(
                    word_table.cell(r_idx, c_idx),
                    txt,
                    bold=bold_flag,
                    center=center_flag,
                    cell_w_cm=cell_w_cm_local,
                )

        # Fill the merged title/page cell once, vertically centered
        # within the full-height column.
        #
        # The title text returned by ``_extract_title_and_page`` is
        # already ``"渑池\n杜氏宗谱"`` so each fragment occupies one
        # person-column inside the merged cell.  We insert TWO empty
        # person-columns (i.e. ``\n\n\n``) between every pair of
        # logical labels so that, vertically, there is roughly a
        # two-character gap between "渑池" and "杜氏宗谱" and between
        # "杜氏宗谱" and the page number "三〇八", matching the
        # spacing of the original document.
        if has_title_col:
            label_groups: List[str] = []
            if title_text:
                # Split the multi-fragment title into individual
                # labels and insert two empty person-columns between
                # consecutive labels.
                title_labels = [s for s in title_text.split("\n") if s]
                if title_labels:
                    label_groups.append("\n\n\n".join(title_labels))
            if page_text:
                label_groups.append(page_text)
            merged_text = "\n\n\n".join(label_groups)
            set_cell_text(
                word_table.cell(0, title_col_idx),
                merged_text,
                bold=True,
                center=True,
                cell_w_cm=cell_w_cm_local,
            )
        return True

    # ---- Fallback: generic grid rendered into the fixed 5x24 table -----
    word_table = _create_fixed_table(doc)
    for r_idx in range(TABLE_ROWS):
        for c_idx in range(TABLE_COLS):
            txt = ""
            if r_idx < rows and c_idx < cols:
                indices = bucket[r_idx][c_idx] if c_idx < len(bucket[r_idx]) else []
                lines = split_into_vertical_lines(indices, t_boxes, t_texts)
                txt = "\n".join(lines)
            set_cell_text(word_table.cell(r_idx, c_idx), txt)
    return True


# ---------------------------------------------------------------------------
# Document building
# ---------------------------------------------------------------------------

def build_docx(parsing_results: List[Dict[str, Any]], output_path: str) -> None:
    """Build a Word document from the parsed JSON results."""
    doc = Document()
    _setup_a4_page(doc)

    # Shrink the document's Normal style so the implicit paragraph that
    # follows the table contributes the smallest possible vertical space.
    # Without this, Word's default 11pt + 1.15 line spacing + 8pt space
    # after can push the page total just past A4, forcing a 2nd page.
    try:
        normal = doc.styles["Normal"]
        normal.font.size = Pt(1)
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after  = Pt(0)
        normal.paragraph_format.line_spacing = 1.0
    except KeyError:
        pass

    # Strip the default empty paragraph python-docx inserts so the table
    # sits at the very top of the page.
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)

    # Title / page number are extracted **per page** below so that each
    # rendered page carries its own page number ("三〇六", "三〇七", ...)
    # rather than reusing the first one detected in the document.

    for page_idx, page in enumerate(parsing_results):
        pruned = page.get("prunedResult", page)
        parsing_list = pruned.get("parsing_res_list", [])
        table_res_list = pruned.get("table_res_list", [])

        # Title / page number for THIS page only.
        title_text, page_text = _extract_title_and_page_for_page(page)

        # ``table_res_list`` is *per-page*, so the cursor that walks
        # through it must also be per-page.  Resetting it inside the
        # loop ensures the table on page 2/3/4 is rendered just like
        # the one on page 1, instead of falling back to the raw HTML
        # branch (which Word renders as a faint dotted grid with no
        # text -- the very symptom the user reported).
        table_index = 0

        # ``overall_ocr_res`` uses the same global image coordinates as
        # ``cell_box_list``; ``table_ocr_pred`` may use a different (often
        # rotated) local coordinate system, so we prefer the overall result.
        overall = pruned.get("overall_ocr_res", {})
        overall_texts: List[str] = list(overall.get("rec_texts", []))
        overall_boxes: List[Box] = [tuple(b) for b in overall.get("rec_boxes", [])]

        for block in parsing_list:
            label = block.get("block_label", "")
            if label == "table":
                if table_index >= len(table_res_list):
                    doc.add_paragraph(block.get("block_content", ""))
                    continue

                tres = table_res_list[table_index]
                table_index += 1

                cell_box_list = [tuple(b) for b in tres.get("cell_box_list", [])]
                # Some pages include OCR cells from the right-hand
                # page-edge strip ("渑池 / 杜氏宗谱 / page number") in
                # the table cell list.  Those cells span the full page
                # height and confuse the row clustering, so strip them
                # out before building the grid.
                cell_box_list = _filter_main_table_cells(cell_box_list)

                # Filter overall OCR boxes to those overlapping this table
                # region (defined by the union of its cell boxes).
                if cell_box_list:
                    tx1 = min(b[0] for b in cell_box_list)
                    ty1 = min(b[1] for b in cell_box_list)
                    tx2 = max(b[2] for b in cell_box_list)
                    ty2 = max(b[3] for b in cell_box_list)
                    table_region: Box = (tx1, ty1, tx2, ty2)
                else:
                    table_region = (0.0, 0.0, 0.0, 0.0)

                t_texts: List[str] = []
                t_boxes: List[Box] = []
                for txt, bx in zip(overall_texts, overall_boxes):
                    if box_overlap_area(bx, table_region) > 0:
                        t_texts.append(txt)
                        t_boxes.append(bx)

                if not render_table(
                    doc, cell_box_list, t_texts, t_boxes,
                    title_text=title_text, page_text=page_text,
                ):
                    doc.add_paragraph(block.get("block_content", ""))
            # Non-table blocks are intentionally skipped: the user only
            # wants the genealogy table on the page.
            else:
                continue

        if page_idx < len(parsing_results) - 1:
            doc.add_page_break()

    doc.save(output_path)


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main(argv: List[str]) -> int:
    if len(argv) < 2:
        print(__doc__)
        return 1

    input_path = argv[1]
    if len(argv) >= 3:
        output_path = argv[2]
    else:
        base, _ = os.path.splitext(input_path)
        output_path = base + ".docx"

    if not os.path.isfile(input_path):
        print(f"Input file not found: {input_path}")
        return 2

    with open(input_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    # PP-StructureV3 has two output shapes:
    #   * legacy: a JSON list, where each element corresponds to one
    #     parsed page and already exposes ``prunedResult`` directly.
    #   * new   : a JSON object with a top-level ``layoutParsingResults``
    #     array; each element wraps the page in
    #     ``{"prunedResult": ..., "markdown": ..., ...}``.
    # Normalise both shapes to a list of page dicts so the rest of
    # the pipeline does not need to care.
    if isinstance(data, dict):
        if "layoutParsingResults" in data:
            data = data["layoutParsingResults"]
        else:
            data = [data]

    build_docx(data, output_path)
    print(f"Word document generated: {output_path}")
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv))
