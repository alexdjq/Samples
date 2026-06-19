# -*- coding: utf-8 -*-
from docx import Document
from docx.oxml.ns import qn

DOCX = r"C:\Users\alexdu\Downloads\vFlat\家谱\窑湾_关键边框.docx"
doc = Document(DOCX)

print(f"Total tables: {len(doc.tables)}")
for ti, t in enumerate(doc.tables):
    nrows, ncols = len(t.rows), len(t.columns)
    print(f"\n=== Table[{ti}]: {nrows} rows x {ncols} cols ===")
    for r in range(nrows):
        for c in range(ncols):
            txt = t.cell(r, c).text.strip()
            if txt:
                # Replace newlines with | for compact display.
                compact = txt.replace("\n", "|")
                print(f"  [{r}][{c}]: {compact}")
