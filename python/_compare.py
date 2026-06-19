# -*- coding: utf-8 -*-
"""Compare each Word table with each OCR layoutParsingResults page to map them."""
import json
from docx import Document

DOCX = r"C:\Users\alexdu\Downloads\vFlat\家谱\窑湾_关键边框v2.docx"
JSON_PATH = r"C:\Users\alexdu\Downloads\vFlat\家谱\窑湾.json"

with open(JSON_PATH, "r", encoding="utf-8") as fh:
    data = json.load(fh)
results = data["layoutParsingResults"]

doc = Document(DOCX)

print(f"Word tables: {len(doc.tables)}    JSON pages: {len(results)}")
print()

# For each table extract the cell with the most distinctive person name.
for ti, t in enumerate(doc.tables[:8]):
    cells_text = []
    for r in range(len(t.rows)):
        for c in range(len(t.columns)):
            txt = t.cell(r, c).text.strip()
            if txt and len(txt.replace("\n", "")) >= 3:
                cells_text.append(txt.replace("\n", ""))
    sample = "; ".join(cells_text[:4])
    print(f"Table[{ti}]: sample = {sample}")

print()
for pi, page in enumerate(results[:8]):
    try:
        texts = page["prunedResult"]["overall_ocr_res"]["rec_texts"]
    except KeyError:
        print(f"Page[{pi}]: no OCR data")
        continue
    long_texts = [t for t in texts if len(t) >= 3][:4]
    sample = "; ".join(long_texts)
    print(f"Page[{pi}]: sample = {sample}")
