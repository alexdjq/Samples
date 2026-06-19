# -*- coding: utf-8 -*-
"""Verify Table[1] (which actually maps to PDF page 2) now contains
廷贵 / 廷璋 / 廷献 in the 十世 row, and that no other tables regressed."""
from docx import Document

DOCX = r"C:\Users\alexdu\Downloads\vFlat\家谱\窑湾_关键边框v3.docx"
doc = Document(DOCX)

print(f"Total tables: {len(doc.tables)}")

# Verify Table[1] (PDF page 2) -- 十世 row should now be populated.
t = doc.tables[1]
print(f"\nTable[1] (PDF page 2): {len(t.rows)} rows x {len(t.columns)} cols")
last_row = len(t.rows) - 1
for c in range(len(t.columns)):
    txt = t.cell(last_row, c).text.strip()
    if txt:
        print(f"  [last row][{c}]: {txt.replace(chr(10), '|')}")

# Verify Table[2] (PDF page 3) is unaffected.
t = doc.tables[2]
print(f"\nTable[2] (PDF page 3): {len(t.rows)} rows x {len(t.columns)} cols")
last_row = len(t.rows) - 1
for c in range(len(t.columns)):
    txt = t.cell(last_row, c).text.strip()
    if txt:
        print(f"  [last row][{c}]: {txt.replace(chr(10), '|')}")

# Quick check for 廷贵, 廷璋, 廷献 anywhere in Table[1].
print("\nSearch 廷贵 / 廷璋 / 廷献 / 廷龙 / 廷士 in Table[1]:")
t = doc.tables[1]
hits = []
for r in range(len(t.rows)):
    for c in range(len(t.columns)):
        txt = t.cell(r, c).text
        for name in ("廷贵", "廷璋", "廷献", "廷龙", "廷士", "廷臣", "廷彦", "廷召"):
            if name in txt:
                hits.append((name, r, c))
for name, r, c in hits:
    print(f"  {name}  -> [{r}][{c}]")
